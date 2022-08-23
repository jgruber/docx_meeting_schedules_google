import os
import re

from docx import Document
from docx.oxml.text.run import CT_R

dir_path = os.path.dirname(os.path.realpath(__file__))
docx_template_file = f"{dir_path}/s_140_google_template.docx"

# intialize template variables
template_variables = {
    "w1_meeting_date": '',
    "w1_bible_reading": '',
    "w1_chairman": '',
    "w1_second_school_chair": '',
    "w1_meeting_start_time": '',
    "w1_opening_song_numer": '',
    "w1_opening_song_title": '',
    "w1_opening_song_scripture": '',
    "w1_opening_prayer": '',
    "w1_intro_start_time": '',
    "w1_treasures_start_time": '',
    "w1_treasures_description": '',
    "w1_treasures_min": '',
    "w1_treasures_assigned": '',
    "w1_digging_start_time": '',
    "w1_digging_min": '',
    "w1_digging_assigned": '',
    "w1_bible_reading_start_time": '',
    "w1_bible_reading_assignment": '',
    "w1_bible_reading_min": '',
    "w1_bible_reading_counsel_point_index": '',
    "w1_bible_reading_counsel_point_number": '',
    "w1_bible_reading_counsel_point_title": '',
    "w1_bible_reading_assigned_main_aud": '',
    "w1_bible_reading_assigned_second_school": '',
    "w1_apply_1_start_time": '',
    "w1_apply_1_description": '',
    "w1_apply_1_counsel_point_index": '',
    "w1_apply_1_counsel_point_title": '',
    "w1_apply_1_assigned_main_aud": '',
    "w1_apply_1_assistant_main_aud": '',
    "w1_apply_1_assigned_sec_aud": '',
    "w1_apply_1_assistant_sec_aud": '',
    "w1_apply_2_start_time": '',
    "w1_apply_2_description": '',
    "w1_apply_2_counsel_point_index": '',
    "w1_apply_2_counsel_point_title": '',
    "w1_apply_2_assigned_main_aud": '',
    "w1_apply_2_assistant_main_aud": '',
    "w1_apply_2_assigned_sec_aud": '',
    "w1_apply_2_assistant_sec_aud": '',
    "w1_apply_3_start_time": '',
    "w1_apply_3_description": '',
    "w1_apply_3_counsel_point_index": '',
    "w1_apply_3_counsel_point_title": '',
    "w1_apply_3_assigned_main_aud": '',
    "w1_apply_3_assistant_main_aud": '',
    "w1_apply_3_assigned_sec_aud": '',
    "w1_apply_3_assistant_sec_aud": '',
    "w1_middle_song_start_time": '',
    "w1_middle_song_numer": '',
    "w1_middle_song_title": '',
    "w1_middle_song_scripture": '',
    "w1_living_1_start_time": '',
    "w1_living_1_description": '',
    "w1_living_1_assigned": '',
    "w1_living_2_start_time": '',
    "w1_living_2_description": '',
    "w1_living_2_assigned": '',
    "w1_cbs_start_time": '',
    "w1_cbs_description": '',
    "w1_cbs_conductor": '',
    "w1_cbs_reader": '',
    "w1_cc_start_time": '',
    "w1_closing_song_start_time": '',
    "w1_closing_song_numer": '',
    "w1_closing_song_title": '',
    "w1_closing_song_scripture": '',
    "w1_closing_prayer": '',
    "w2_meeting_date": '',
    "w2_bible_reading": '',
    "w2_chairman": '',
    "w2_second_school_chair": '',
    "w2_meeting_start_time": '',
    "w2_opening_song_numer": '',
    "w2_opening_song_title": '',
    "w2_opening_song_scripture": '',
    "w2_opening_prayer": '',
    "w2_intro_start_time": '',
    "w2_treasures_start_time": '',
    "w2_treasures_description": '',
    "w2_treasures_min": '',
    "w2_treasures_assigned": '',
    "w2_digging_start_time": '',
    "w2_digging_min": '',
    "w2_digging_assigned": '',
    "w2_bible_reading_start_time": '',
    "w2_bible_reading_assignment": '',
    "w2_bible_reading_min": '',
    "w2_bible_reading_counsel_point_index": '',
    "w2_bible_reading_counsel_point_number": '',
    "w2_bible_reading_counsel_point_title": '',
    "w2_bible_reading_assigned_main_aud": '',
    "w2_bible_reading_assigned_second_school": '',
    "w2_apply_1_start_time": '',
    "w2_apply_1_description": '',
    "w2_apply_1_counsel_point_index": '',
    "w2_apply_1_counsel_point_title": '',
    "w2_apply_1_assigned_main_aud": '',
    "w2_apply_1_assistant_main_aud": '',
    "w2_apply_1_assigned_sec_aud": '',
    "w2_apply_1_assistant_sec_aud": '',
    "w2_apply_2_start_time": '',
    "w2_apply_2_description": '',
    "w2_apply_2_counsel_point_index": '',
    "w2_apply_2_counsel_point_title": '',
    "w2_apply_2_assigned_main_aud": '',
    "w2_apply_2_assistant_main_aud": '',
    "w2_apply_2_assigned_sec_aud": '',
    "w2_apply_2_assistant_sec_aud": '',
    "w2_apply_3_start_time": '',
    "w2_apply_3_description": '',
    "w2_apply_3_counsel_point_index": '',
    "w2_apply_3_counsel_point_title": '',
    "w2_apply_3_assigned_main_aud": '',
    "w2_apply_3_assistant_main_aud": '',
    "w2_apply_3_assigned_sec_aud": '',
    "w2_apply_3_assistant_sec_aud": '',
    "w2_middle_song_start_time": '',
    "w2_middle_song_numer": '',
    "w2_middle_song_title": '',
    "w2_middle_song_scripture": '',
    "w2_living_1_start_time": '',
    "w2_living_1_description": '',
    "w2_living_1_assigned": '',
    "w2_living_2_start_time": '',
    "w2_living_2_description": '',
    "w2_living_2_assigned": '',
    "w2_cbs_start_time": '',
    "w2_cbs_description": '',
    "w2_cbs_conductor": '',
    "w2_cbs_reader": '',
    "w2_cc_start_time": '',
    "w2_closing_song_start_time": '',
    "w2_closing_song_numer": '',
    "w2_closing_song_title": '',
    "w2_closing_song_scripture": '',
    "w2_closing_prayer": '',
    "w3_meeting_date": '',
    "w3_bible_reading": '',
    "w3_chairman": '',
    "w3_second_school_chair": '',
    "w3_meeting_start_time": '',
    "w3_opening_song_numer": '',
    "w3_opening_song_title": '',
    "w3_opening_song_scripture": '',
    "w3_opening_prayer": '',
    "w3_intro_start_time": '',
    "w3_treasures_start_time": '',
    "w3_treasures_description": '',
    "w3_treasures_min": '',
    "w3_treasures_assigned": '',
    "w3_digging_start_time": '',
    "w3_digging_min": '',
    "w3_digging_assigned": '',
    "w3_bible_reading_start_time": '',
    "w3_bible_reading_assignment": '',
    "w3_bible_reading_min": '',
    "w3_bible_reading_counsel_point_index": '',
    "w3_bible_reading_counsel_point_number": '',
    "w3_bible_reading_counsel_point_title": '',
    "w3_bible_reading_assigned_main_aud": '',
    "w3_bible_reading_assigned_second_school": '',
    "w3_apply_1_start_time": '',
    "w3_apply_1_description": '',
    "w3_apply_1_counsel_point_index": '',
    "w3_apply_1_counsel_point_title": '',
    "w3_apply_1_assigned_main_aud": '',
    "w3_apply_1_assistant_main_aud": '',
    "w3_apply_1_assigned_sec_aud": '',
    "w3_apply_1_assistant_sec_aud": '',
    "w3_apply_2_start_time": '',
    "w3_apply_2_description": '',
    "w3_apply_2_counsel_point_index": '',
    "w3_apply_2_counsel_point_title": '',
    "w3_apply_2_assigned_main_aud": '',
    "w3_apply_2_assistant_main_aud": '',
    "w3_apply_2_assigned_sec_aud": '',
    "w3_apply_2_assistant_sec_aud": '',
    "w3_apply_3_start_time": '',
    "w3_apply_3_description": '',
    "w3_apply_3_counsel_point_index": '',
    "w3_apply_3_counsel_point_title": '',
    "w3_apply_3_assigned_main_aud": '',
    "w3_apply_3_assistant_main_aud": '',
    "w3_apply_3_assigned_sec_aud": '',
    "w3_apply_3_assistant_sec_aud": '',
    "w3_middle_song_start_time": '',
    "w3_middle_song_numer": '',
    "w3_middle_song_title": '',
    "w3_middle_song_scripture": '',
    "w3_living_1_start_time": '',
    "w3_living_1_description": '',
    "w3_living_1_assigned": '',
    "w3_living_2_start_time": '',
    "w3_living_2_description": '',
    "w3_living_2_assigned": '',
    "w3_cbs_start_time": '',
    "w3_cbs_description": '',
    "w3_cbs_conductor": '',
    "w3_cbs_reader": '',
    "w3_cc_start_time": '',
    "w3_closing_song_start_time": '',
    "w3_closing_song_numer": '',
    "w3_closing_song_title": '',
    "w3_closing_song_scripture": '',
    "w3_closing_prayer": '',
    "w4_meeting_date": '',
    "w4_bible_reading": '',
    "w4_chairman": '',
    "w4_second_school_chair": '',
    "w4_meeting_start_time": '',
    "w4_opening_song_numer": '',
    "w4_opening_song_title": '',
    "w4_opening_song_scripture": '',
    "w4_opening_prayer": '',
    "w4_intro_start_time": '',
    "w4_treasures_start_time": '',
    "w4_treasures_description": '',
    "w4_treasures_min": '',
    "w4_treasures_assigned": '',
    "w4_digging_start_time": '',
    "w4_digging_min": '',
    "w4_digging_assigned": '',
    "w4_bible_reading_start_time": '',
    "w4_bible_reading_assignment": '',
    "w4_bible_reading_min": '',
    "w4_bible_reading_counsel_point_index": '',
    "w4_bible_reading_counsel_point_number": '',
    "w4_bible_reading_counsel_point_title": '',
    "w4_bible_reading_assigned_main_aud": '',
    "w4_bible_reading_assigned_second_school": '',
    "w4_apply_1_start_time": '',
    "w4_apply_1_description": '',
    "w4_apply_1_counsel_point_index": '',
    "w4_apply_1_counsel_point_title": '',
    "w4_apply_1_assigned_main_aud": '',
    "w4_apply_1_assistant_main_aud": '',
    "w4_apply_1_assigned_sec_aud": '',
    "w4_apply_1_assistant_sec_aud": '',
    "w4_apply_2_start_time": '',
    "w4_apply_2_description": '',
    "w4_apply_2_counsel_point_index": '',
    "w4_apply_2_counsel_point_title": '',
    "w4_apply_2_assigned_main_aud": '',
    "w4_apply_2_assistant_main_aud": '',
    "w4_apply_2_assigned_sec_aud": '',
    "w4_apply_2_assistant_sec_aud": '',
    "w4_apply_3_start_time": '',
    "w4_apply_3_description": '',
    "w4_apply_3_counsel_point_index": '',
    "w4_apply_3_counsel_point_title": '',
    "w4_apply_3_assigned_main_aud": '',
    "w4_apply_3_assistant_main_aud": '',
    "w4_apply_3_assigned_sec_aud": '',
    "w4_apply_3_assistant_sec_aud": '',
    "w4_middle_song_start_time": '',
    "w4_middle_song_numer": '',
    "w4_middle_song_title": '',
    "w4_middle_song_scripture": '',
    "w4_living_1_start_time": '',
    "w4_living_1_description": '',
    "w4_living_1_assigned": '',
    "w4_living_2_start_time": '',
    "w4_living_2_description": '',
    "w4_living_2_assigned": '',
    "w4_cbs_start_time": '',
    "w4_cbs_description": '',
    "w4_cbs_conductor": '',
    "w4_cbs_reader": '',
    "w4_cc_start_time": '',
    "w4_closing_song_start_time": '',
    "w4_closing_song_numer": '',
    "w4_closing_song_title": '',
    "w4_closing_song_scripture": '',
    "w4_closing_prayer": '',
    "w5_meeting_date": '',
    "w5_bible_reading": '',
    "w5_chairman": '',
    "w5_second_school_chair": '',
    "w5_meeting_start_time": '',
    "w5_opening_song_numer": '',
    "w5_opening_song_title": '',
    "w5_opening_song_scripture": '',
    "w5_opening_prayer": '',
    "w5_intro_start_time": '',
    "w5_treasures_start_time": '',
    "w5_treasures_description": '',
    "w5_treasures_min": '',
    "w5_treasures_assigned": '',
    "w5_digging_start_time": '',
    "w5_digging_min": '',
    "w5_digging_assigned": '',
    "w5_bible_reading_start_time": '',
    "w5_bible_reading_assignment": '',
    "w5_bible_reading_min": '',
    "w5_bible_reading_counsel_point_index": '',
    "w5_bible_reading_counsel_point_number": '',
    "w5_bible_reading_counsel_point_title": '',
    "w5_bible_reading_assigned_main_aud": '',
    "w5_bible_reading_assigned_second_school": '',
    "w5_apply_1_start_time": '',
    "w5_apply_1_description": '',
    "w5_apply_1_counsel_point_index": '',
    "w5_apply_1_counsel_point_title": '',
    "w5_apply_1_assigned_main_aud": '',
    "w5_apply_1_assistant_main_aud": '',
    "w5_apply_1_assigned_sec_aud": '',
    "w5_apply_1_assistant_sec_aud": '',
    "w5_apply_2_start_time": '',
    "w5_apply_2_description": '',
    "w5_apply_2_counsel_point_index": '',
    "w5_apply_2_counsel_point_title": '',
    "w5_apply_2_assigned_main_aud": '',
    "w5_apply_2_assistant_main_aud": '',
    "w5_apply_2_assigned_sec_aud": '',
    "w5_apply_2_assistant_sec_aud": '',
    "w5_apply_3_start_time": '',
    "w5_apply_3_description": '',
    "w5_apply_3_counsel_point_index": '',
    "w5_apply_3_counsel_point_title": '',
    "w5_apply_3_assigned_main_aud": '',
    "w5_apply_3_assistant_main_aud": '',
    "w5_apply_3_assigned_sec_aud": '',
    "w5_apply_3_assistant_sec_aud": '',
    "w5_middle_song_start_time": '',
    "w5_middle_song_numer": '',
    "w5_middle_song_title": '',
    "w5_middle_song_scripture": '',
    "w5_living_1_start_time": '',
    "w5_living_1_description": '',
    "w5_living_1_assigned": '',
    "w5_living_2_start_time": '',
    "w5_living_2_description": '',
    "w5_living_2_assigned": '',
    "w5_cbs_start_time": '',
    "w5_cbs_description": '',
    "w5_cbs_conductor": '',
    "w5_cbs_reader": '',
    "w5_cc_start_time": '',
    "w5_closing_song_start_time": '',
    "w5_closing_song_numer": '',
    "w5_closing_song_title": '',
    "w5_closing_song_scripture": '',
    "w5_closing_prayer": ''
}

# give bounds to template tables for each week
template_table_indexes = {
    "w1_start": 0,
    "w1_end": 7,
    "w2_start": 8,
    "w2_end": 15,
    "w3_start": 16,
    "w3_end": 23,
    "w4_start": 24,
    "w4_end": 31,
    "w5_start": 32,
    "w5_end": 39,
}

def render(meetings=[], docx_filename=docx_template_file):
    variables = dict(template_variables)
    number_of_meetings = len(meetings)
    for i in range(0, number_of_meetings):
        variables[f"w{i+1}_meeting_date"] = meetings[i]['meeting_date']
        variables[f"w{i+1}_bible_reading"] = meetings[i]['bible_reading']
        variables[f"w{i+1}_chairman"] = ''
        variables[f"w{i+1}_second_school_chair"] = ''
        variables[f"w{i+1}_meeting_start_time"] = meetings[i]['opening_song']['start']
        variables[f"w{i+1}_opening_song_number"] = meetings[i]['opening_song']['details']['number']
        variables[f"w{i+1}_opening_song_title"] = meetings[i]['opening_song']['details']['title']
        variables[f"w{i+1}_opening_song_scripture"] = meetings[i]['opening_song']['details']['ref_scripture']
        variables[f"w{i+1}_opening_prayer"] = ''
        variables[f"w{i+1}_intro_start_time"] = meetings[i]['opening_comments']['start']
        variables[f"w{i+1}_treasures_start_time"] = meetings[i]['treasures_from_gods_word']['parts'][0]['start']
        variables[f"w{i+1}_treasures_description"] = meetings[i]['treasures_from_gods_word']['parts'][0]['theme']
        variables[f"w{i+1}_treasures_min"] = meetings[i]['treasures_from_gods_word']['parts'][0]['duration_min']
        variables[f"w{i+1}_treasures_assigned"] = ''
        variables[f"w{i+1}_digging_start_time"] = meetings[i]['treasures_from_gods_word']['parts'][1]['start']
        variables[f"w{i+1}_digging_min"] = meetings[i]['treasures_from_gods_word']['parts'][1]['duration_min']
        variables[f"w{i+1}_digging_assigned"] = ''
        variables[f"w{i+1}_bible_reading_start_time"] = meetings[i]['treasures_from_gods_word']['parts'][2]['start']
        description = meetings[i]['treasures_from_gods_word']['parts'][2]['reading']
        variables[f"w{i+1}_bible_reading_assignment"] = re.sub(r'\(th study [123456789][0123456789]?\)', '', description)
        variables[f"w{i+1}_bible_reading_min"] = meetings[i]['treasures_from_gods_word']['parts'][2]['duration_min']
        variables[f"w{i+1}_bible_reading_counsel_point_index"] = meetings[i]['treasures_from_gods_word']['parts'][2]['counsel_point']['index']
        variables[f"w{i+1}_bible_reading_counsel_point_number"] = meetings[i]['treasures_from_gods_word']['parts'][2]['counsel_point']['number']
        variables[f"w{i+1}_bible_reading_counsel_point_title"] = meetings[i]['treasures_from_gods_word']['parts'][2]['counsel_point']['title']
        variables[f"w{i+1}_bible_reading_assigned_main_aud"] = ''
        variables[f"w{i+1}_bible_reading_assigned_second_school"] = ''
        apply_part_count = len(meetings[i]['apply_yourself_to_the_field_ministry']['parts'])
        if apply_part_count > 0:
            variables[f"w{i+1}_apply_1_start_time"] = meetings[i]['apply_yourself_to_the_field_ministry']['parts'][0]['start']
            description = meetings[i]['apply_yourself_to_the_field_ministry']['parts'][0]['description']
            variables[f"w{i+1}_apply_1_description"] = re.sub(r'\(th study [123456789][0123456789]?\)', '', description)
            if 'number' in meetings[i]['apply_yourself_to_the_field_ministry']['parts'][0]['counsel_point']:
                variables[f"w{i+1}_apply_1_counsel_point_index"] = meetings[i]['apply_yourself_to_the_field_ministry']['parts'][0]['counsel_point']['index']
                variables[f"w{i+1}_apply_1_counsel_point_number"] = meetings[i]['apply_yourself_to_the_field_ministry']['parts'][0]['counsel_point']['number']
                variables[f"w{i+1}_apply_1_counsel_point_title"] = meetings[i]['apply_yourself_to_the_field_ministry']['parts'][0]['counsel_point']['title']
            else:
                variables[f"w{i+1}_apply_1_counsel_point_index"] = ''
                variables[f"w{i+1}_apply_1_counsel_point_number"] = ''
                variables[f"w{i+1}_apply_1_counsel_point_title"] = ''
            variables[f"w{i+1}_apply_1_assigned_main_aud"] = ''
            variables[f"w{i+1}_apply_1_assistant_main_aud"] = ''
            variables[f"w{i+1}_apply_1_assigned_sec_aud"] = ''
            variables[f"w{i+1}_apply_1_assistant_sec_aud"] = ''
        else:
            variables[f"w{i+1}_apply_1_start_time"] = ''
            variables[f"w{i+1}_apply_1_description"] = ''
            variables[f"w{i+1}_apply_1_counsel_point_index"] = ''
            variables[f"w{i+1}_apply_1_counsel_point_number"] = ''
            variables[f"w{i+1}_apply_1_counsel_point_title"] = ''
            variables[f"w{i+1}_apply_1_assigned_main_aud"] = ''
            variables[f"w{i+1}_apply_1_assistant_main_aud"] = ''
            variables[f"w{i+1}_apply_1_assigned_sec_aud"] = ''
            variables[f"w{i+1}_apply_1_assistant_sec_aud"] = ''
        if apply_part_count > 1:
            variables[f"w{i+1}_apply_2_start_time"] = meetings[i]['apply_yourself_to_the_field_ministry']['parts'][1]['start']
            description = meetings[i]['apply_yourself_to_the_field_ministry']['parts'][1]['description']
            variables[f"w{i+1}_apply_2_description"] = re.sub(r'\(th study [123456789][0123456789]?\)', '', description)
            if 'number' in meetings[i]['apply_yourself_to_the_field_ministry']['parts'][1]['counsel_point']:
                variables[f"w{i+1}_apply_2_counsel_point_index"] = meetings[i]['apply_yourself_to_the_field_ministry']['parts'][1]['counsel_point']['index']
                variables[f"w{i+1}_apply_2_counsel_point_number"] = meetings[i]['apply_yourself_to_the_field_ministry']['parts'][1]['counsel_point']['number']
                variables[f"w{i+1}_apply_2_counsel_point_title"] = meetings[i]['apply_yourself_to_the_field_ministry']['parts'][1]['counsel_point']['title']
            else:
                variables[f"w{i+1}_apply_2_counsel_point_index"] = ''
                variables[f"w{i+1}_apply_2_counsel_point_number"] = ''
                variables[f"w{i+1}_apply_2_counsel_point_title"] = ''
            variables[f"w{i+1}_apply_2_assigned_main_aud"] = ''
            variables[f"w{i+1}_apply_2_assistant_main_aud"] = ''
            variables[f"w{i+1}_apply_2_assigned_sec_aud"] = ''
            variables[f"w{i+1}_apply_2_assistant_sec_aud"] = ''
        else:
            variables[f"w{i+1}_apply_2_start_time"] = ''
            variables[f"w{i+1}_apply_2_description"] = ''
            variables[f"w{i+1}_apply_2_counsel_point_index"] = ''
            variables[f"w{i+1}_apply_2_counsel_point_number"] = ''
            variables[f"w{i+1}_apply_2_counsel_point_title"] = ''
            variables[f"w{i+1}_apply_2_assigned_main_aud"] = ''
            variables[f"w{i+1}_apply_2_assistant_main_aud"] = ''
            variables[f"w{i+1}_apply_2_assigned_sec_aud"] = ''
            variables[f"w{i+1}_apply_2_assistant_sec_aud"] = ''
        if apply_part_count > 2:
            variables[f"w{i+1}_apply_3_start_time"] = meetings[i]['apply_yourself_to_the_field_ministry']['parts'][2]['start']
            description = meetings[i]['apply_yourself_to_the_field_ministry']['parts'][2]['description']
            variables[f"w{i+1}_apply_3_description"] = re.sub(r'\(th study [123456789][0123456789]?\)', '', description)
            if 'number' in meetings[i]['apply_yourself_to_the_field_ministry']['parts'][2]['counsel_point']:
                variables[f"w{i+1}_apply_3_counsel_point_index"] = meetings[i]['apply_yourself_to_the_field_ministry']['parts'][2]['counsel_point']['index']
                variables[f"w{i+1}_apply_3_counsel_point_number"] = meetings[i]['apply_yourself_to_the_field_ministry']['parts'][2]['counsel_point']['number']
                variables[f"w{i+1}_apply_3_counsel_point_title"] = meetings[i]['apply_yourself_to_the_field_ministry']['parts'][2]['counsel_point']['title']
            else:
                variables[f"w{i+1}_apply_3_counsel_point_index"] = ''
                variables[f"w{i+1}_apply_3_counsel_point_number"] = ''
                variables[f"w{i+1}_apply_3_counsel_point_title"] = ''
            variables[f"w{i+1}_apply_3_assigned_main_aud"] = ''
            variables[f"w{i+1}_apply_3_assistant_main_aud"] = ''
            variables[f"w{i+1}_apply_3_assigned_sec_aud"] = ''
            variables[f"w{i+1}_apply_3_assistant_sec_aud"] = ''
        else:
            variables[f"w{i+1}_apply_3_start_time"] = ''
            variables[f"w{i+1}_apply_3_description"] = ''
            variables[f"w{i+1}_apply_3_counsel_point_index"] = ''
            variables[f"w{i+1}_apply_3_counsel_point_number"] = ''
            variables[f"w{i+1}_apply_3_counsel_point_title"] = ''
            variables[f"w{i+1}_apply_3_assigned_main_aud"] = ''
            variables[f"w{i+1}_apply_3_assistant_main_aud"] = ''
            variables[f"w{i+1}_apply_3_assigned_sec_aud"] = ''
            variables[f"w{i+1}_apply_3_assistant_sec_aud"] = ''
        variables[f"w{i+1}_middle_song_start_time"] = meetings[i]['middle_song']['start']
        variables[f"w{i+1}_middle_song_number"] = meetings[i]['middle_song']['details']['number']
        variables[f"w{i+1}_middle_song_title"] = meetings[i]['middle_song']['details']['title']
        variables[f"w{i+1}_middle_song_scripture"] = meetings[i]['middle_song']['details']['ref_scripture']
        living_parts = len(meetings[i]['living_as_christians']['parts']) - 1
        if living_parts > 0:
            variables[f"w{i+1}_living_1_start_time"] = meetings[i]['living_as_christians']['parts'][0]['start']
            variables[f"w{i+1}_living_1_description"] = meetings[i]['living_as_christians']['parts'][0]['description']
            variables[f"w{i+1}_living_1_assigned"] = ''
        else:
            variables[f"w{i+1}_living_1_start_time"] = ''
            variables[f"w{i+1}_living_1_description"] = ''
            variables[f"w{i+1}_living_1_assigned"] = ''
        if living_parts > 1:
            variables[f"w{i+1}_living_2_start_time"] = meetings[i]['living_as_christians']['parts'][1]['start']
            variables[f"w{i+1}_living_2_description"] = meetings[i]['living_as_christians']['parts'][1]['description']
            variables[f"w{i+1}_living_2_assigned"] = ''
        else:
            variables[f"w{i+1}_living_2_start_time"] = ''
            variables[f"w{i+1}_living_2_description"] = ''
            variables[f"w{i+1}_living_2_assigned"] = ''
        variables[f"w{i+1}_cbs_start_time"] = meetings[i]['living_as_christians']['parts'][-1]['start']
        variables[f"w{i+1}_cbs_description"] = meetings[i]['living_as_christians']['parts'][-1]['description']
        variables[f"w{i+1}_cbs_conductor"] = ''
        variables[f"w{i+1}_cbs_reader"] = ''
        variables[f"w{i+1}_cc_start_time"] = meetings[i]['concluding_comments']['start']
        variables[f"w{i+1}_closing_song_start_time"] = meetings[i]['closing_song']['start']
        variables[f"w{i+1}_closing_song_number"] = meetings[i]['closing_song']['details']['number']
        variables[f"w{i+1}_closing_song_title"] = meetings[i]['closing_song']['details']['title']
        variables[f"w{i+1}_closing_song_scripture"] = meetings[i]['closing_song']['details']['ref_scripture']
        variables[f"w{i+1}_closing_prayer"] = ''
    if number_of_meetings < 4:
        variables["w4_meeting_date"] = ''
        variables["w4_bible_reading"] = ''
        variables["w4_chairman"] = ''
        variables["w4_second_school_chair"] = ''
        variables["w4_meeting_start_time"] = ''
        variables["w4_opening_song_numer"] = ''
        variables["w4_opening_song_title"] = ''
        variables["w4_opening_song_scripture"] = ''
        variables["w4_opening_prayer"] = ''
        variables["w4_intro_start_time"] = ''
        variables["w4_treasures_start_time"] = ''
        variables["w4_treasures_description"] = ''
        variables["w4_treasures_min"] = ''
        variables["w4_treasures_assigned"] = ''
        variables["w4_digging_start_time"] = ''
        variables["w4_digging_min"] = ''
        variables["w4_digging_assigned"] = ''
        variables["w4_bible_reading_start_time"] = ''
        variables["w4_bible_reading_assignment"] = ''
        variables["w4_bible_reading_min"] = ''
        variables["w4_bible_reading_counsel_point_index"] = ''
        variables["w4_bible_reading_counsel_point_number"] = ''
        variables["w4_bible_reading_counsel_point_title"] = ''
        variables["w4_bible_reading_assigned_main_aud"] = ''
        variables["w4_bible_reading_assigned_second_school"] = ''
        variables["w4_apply_1_start_time"] = ''
        variables["w4_apply_1_description"] = ''
        variables["w4_apply_1_counsel_point_index"] = ''
        variables["w4_apply_1_counsel_point_number"] = ''
        variables["w4_apply_1_counsel_point_title"] = ''
        variables["w4_apply_1_assigned_main_aud"] = ''
        variables["w4_apply_1_assistant_main_aud"] = ''
        variables["w4_apply_1_assigned_sec_aud"] = ''
        variables["w4_apply_1_assistant_sec_aud"] = ''
        variables["w4_apply_2_start_time"] = ''
        variables["w4_apply_2_description"] = ''
        variables["w4_apply_2_counsel_point_index"] = ''
        variables["w4_apply_2_counsel_point_number"] = ''
        variables["w4_apply_2_counsel_point_title"] = ''
        variables["w4_apply_2_assigned_main_aud"] = ''
        variables["w4_apply_2_assistant_main_aud"] = ''
        variables["w4_apply_2_assigned_sec_aud"] = ''
        variables["w4_apply_2_assistant_sec_aud"] = ''
        variables["w4_apply_3_start_time"] = ''
        variables["w4_apply_3_description"] = ''
        variables["w4_apply_3_counsel_point_index"] = ''
        variables["w4_apply_3_counsel_point_number"] = ''
        variables["w4_apply_3_counsel_point_title"] = ''
        variables["w4_apply_3_assigned_main_aud"] = ''
        variables["w4_apply_3_assistant_main_aud"] = ''
        variables["w4_apply_3_assigned_sec_aud"] = ''
        variables["w4_apply_3_assistant_sec_aud"] = ''
        variables["w4_middle_song_start_time"] = ''
        variables["w4_middle_song_numer"] = ''
        variables["w4_middle_song_title"] = ''
        variables["w4_middle_song_scripture"] = ''
        variables["w4_living_1_start_time"] = ''
        variables["w4_living_1_description"] = ''
        variables["w4_living_1_assigned"] = ''
        variables["w4_living_2_start_time"] = ''
        variables["w4_living_2_description"] = ''
        variables["w4_living_2_assigned"] = ''
        variables["w4_cbs_start_time"] = ''
        variables["w4_cbs_description"] = ''
        variables["w4_cbs_conductor"] = ''
        variables["w4_cbs_reader"] = ''
        variables["w4_cc_start_time"] = ''
        variables["w4_closing_song_start_time"] = ''
        variables["w4_closing_song_numer"] = ''
        variables["w4_closing_song_title"] = ''
        variables["w4_closing_song_scripture"] = ''
        variables["w4_closing_prayer"] = ''
    if number_of_meetings < 5:
        variables["w5_meeting_date"] = ''
        variables["w5_bible_reading"] = ''
        variables["w5_chairman"] = ''
        variables["w5_second_school_chair"] = ''
        variables["w5_meeting_start_time"] = ''
        variables["w5_opening_song_numer"] = ''
        variables["w5_opening_song_title"] = ''
        variables["w5_opening_song_scripture"] = ''
        variables["w5_opening_prayer"] = ''
        variables["w5_intro_start_time"] = ''
        variables["w5_treasures_start_time"] = ''
        variables["w5_treasures_description"] = ''
        variables["w5_treasures_min"] = ''
        variables["w5_treasures_assigned"] = ''
        variables["w5_digging_start_time"] = ''
        variables["w5_digging_min"] = ''
        variables["w5_digging_assigned"] = ''
        variables["w5_bible_reading_start_time"] = ''
        variables["w5_bible_reading_assignment"] = ''
        variables["w5_bible_reading_min"] = ''
        variables["w5_bible_reading_counsel_point_index"] = ''
        variables["w5_bible_reading_counsel_point_number"] = ''
        variables["w5_bible_reading_counsel_point_title"] = ''
        variables["w5_bible_reading_assigned_main_aud"] = ''
        variables["w5_bible_reading_assigned_second_school"] = ''
        variables["w5_apply_1_start_time"] = ''
        variables["w5_apply_1_description"] = ''
        variables["w5_apply_1_counsel_point_index"] = ''
        variables["w5_apply_1_counsel_point_number"] = ''
        variables["w5_apply_1_counsel_point_title"] = ''
        variables["w5_apply_1_assigned_main_aud"] = ''
        variables["w5_apply_1_assistant_main_aud"] = ''
        variables["w5_apply_1_assigned_sec_aud"] = ''
        variables["w5_apply_1_assistant_sec_aud"] = ''
        variables["w5_apply_2_start_time"] = ''
        variables["w5_apply_2_description"] = ''
        variables["w5_apply_2_counsel_point_index"] = ''
        variables["w5_apply_2_counsel_point_number"] = ''
        variables["w5_apply_2_counsel_point_title"] = ''
        variables["w5_apply_2_assigned_main_aud"] = ''
        variables["w5_apply_2_assistant_main_aud"] = ''
        variables["w5_apply_2_assigned_sec_aud"] = ''
        variables["w5_apply_2_assistant_sec_aud"] = ''
        variables["w5_apply_3_start_time"] = ''
        variables["w5_apply_3_description"] = ''
        variables["w5_apply_3_counsel_point_index"] = ''
        variables["w5_apply_3_counsel_point_number"] = ''
        variables["w5_apply_3_counsel_point_title"] = ''
        variables["w5_apply_3_assigned_main_aud"] = ''
        variables["w5_apply_3_assistant_main_aud"] = ''
        variables["w5_apply_3_assigned_sec_aud"] = ''
        variables["w5_apply_3_assistant_sec_aud"] = ''
        variables["w5_middle_song_start_time"] = ''
        variables["w5_middle_song_numer"] = ''
        variables["w5_middle_song_title"] = ''
        variables["w5_middle_song_scripture"] = ''
        variables["w5_living_1_start_time"] = ''
        variables["w5_living_1_description"] = ''
        variables["w5_living_1_assigned"] = ''
        variables["w5_living_2_start_time"] = ''
        variables["w5_living_2_description"] = ''
        variables["w5_living_2_assigned"] = ''
        variables["w5_cbs_start_time"] = ''
        variables["w5_cbs_description"] = ''
        variables["w5_cbs_conductor"] = ''
        variables["w5_cbs_reader"] = ''
        variables["w5_cc_start_time"] = ''
        variables["w5_closing_song_start_time"] = ''
        variables["w5_closing_song_numer"] = ''
        variables["w5_closing_song_title"] = ''
        variables["w5_closing_song_scripture"] = ''
        variables["w5_closing_prayer"] = ''
    dir_path = os.path.dirname(os.path.realpath(__file__))
    doc = Document(docx_filename)
    var_regex = re.compile("{{.*}}")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        replacement_text = str(r.text)
                        found_variable = False
                        for m in re.finditer(r'\{\{[^\}\}]*\}\}', r.text):
                            replace_str = r.text[m.span()[0]:m.span()[1]]
                            variable_name = r.text[m.span()[0]+3:m.span()[1]-3].strip()
                            variable_value = ''
                            if variable_name in variables:
                                variable_value = variables[variable_name]
                                found_variable = True
                            replacement_text = replacement_text.replace(replace_str, str(variable_value))
                        if found_variable:
                            r.text = replacement_text
    # remove unused template tables
    for i, t in enumerate(doc.tables):
        if number_of_meetings < 4 and i >= template_table_indexes['w4_start']:
            table_element = t._element
            part_element = table_element.getnext()
            body_element = table_element.getparent()
            body_element.remove(table_element)
            body_element.remove(part_element)
        elif number_of_meetings < 5 and i >= template_table_indexes['w5_start']:
            table_element = t._element
            part_element = table_element.getnext()
            body_element = t._element.getparent()
            body_element.remove(table_element)
            body_element.remove(part_element)
    # remove unused page breaks
    if number_of_meetings < 4:
        last_table_element = doc.tables[template_table_indexes['w3_end']]._element
        paragraph_element = last_table_element.getnext()
        for child in paragraph_element.getchildren():
            if isinstance(child, CT_R):
                paragraph_element.remove(child)
    if number_of_meetings < 5:
        last_table_element = doc.tables[template_table_indexes['w4_end']]._element
        paragraph_element = last_table_element.getnext()
        for child in paragraph_element.getchildren():
            if isinstance(child, CT_R):
                paragraph_element.remove(child)
    return doc
