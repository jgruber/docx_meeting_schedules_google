import os
import re

from docx import Document
from docx.oxml.text.run import CT_R

dir_path = os.path.dirname(os.path.realpath(__file__))
docx_template_file = f"{dir_path}/s_140_google_template.docx"

# intialize template variables
template_variables = {
    "w1_week": '',
    "w1_meeting_date": '',
    "w1_opening_song_start": '',
    "w1_opening_song_number": '',
    "w1_opening_song_title": '',
    "w1_opening_song_scripture": '',
    "w1_chairman": '',
    "w1_public_talk_start": '',
    "w1_public_talk_title": '',
    "w1_public_speaker": '',
    "w1_public_speaker_congregation": '',
    "w1_middle_song_start": '',
    "w1_middle_song_number": '',
    "w1_middle_song_title": '',
    "w1_middle_song_scripture": '',
    "w1_watchtower_start": '',
    "w1_watchtower_title": '',
    "w1_watchtower_conductor": '',
    "w1_watchtower_reader": '',
    "w1_closing_song_start": '',
    "w1_closing_song_number": '',
    "w1_closing_song_title": '',
    "w1_closing_song_scripture": '',
    "w1_closing_prayer": '',
    "w1_foyer_attendant": '',
    "w1_auditorium_attendant": '',
    "w1_stage_attendant": '',
    "w1_sound_console_operator": '',
    "w1_video_console_operator": '',
    "w1_mic_handler_1": '',
    "w1_mic_handler_2": '',
    "w1_mic_handler_3": '',
    "w1_mic_handler_4": '',
    "w1_mic_handler_5": '',
    "w1_mic_handler_6": '',
    "w1_hospitality_service_group": '',
    "w1_touch_up_cleaning_service_group": '',
    "w1_outgoing_speaker_1": '',
    "w1_outgoing_speaker_congregation_1": '',
    "w1_outgoing_speaker_2": '',
    "w1_outgoing_speaker_congregation_2": '',
    "w2_week": '',
    "w2_meeting_date": '',
    "w2_opening_song_start": '',
    "w2_opening_song_number": '',
    "w2_opening_song_title": '',
    "w2_opening_song_scripture": '',
    "w2_chairman": '',
    "w2_public_talk_start": '',
    "w2_public_talk_title": '',
    "w2_public_speaker": '',
    "w2_public_speaker_congregation": '',
    "w2_middle_song_start": '',
    "w2_middle_song_number": '',
    "w2_middle_song_title": '',
    "w2_middle_song_scripture": '',
    "w2_watchtower_start": '',
    "w2_watchtower_title": '',
    "w2_watchtower_conductor": '',
    "w2_watchtower_reader": '',
    "w2_closing_song_start": '',
    "w2_closing_song_number": '',
    "w2_closing_song_title": '',
    "w2_closing_song_scripture": '',
    "w2_closing_prayer": '',
    "w2_foyer_attendant": '',
    "w2_auditorium_attendant": '',
    "w2_stage_attendant": '',
    "w2_sound_console_operator": '',
    "w2_video_console_operator": '',
    "w2_mic_handler_1": '',
    "w2_mic_handler_2": '',
    "w2_mic_handler_3": '',
    "w2_mic_handler_4": '',
    "w2_mic_handler_5": '',
    "w2_mic_handler_6": '',
    "w2_hospitality_service_group": '',
    "w2_touch_up_cleaning_service_group": '',
    "w2_outgoing_speaker_1": '',
    "w2_outgoing_speaker_congregation_1": '',
    "w2_outgoing_speaker_2": '',
    "w2_outgoing_speaker_congregation_2": '',
    "w3_week": '',
    "w3_meeting_date": '',
    "w3_opening_song_start": '',
    "w3_opening_song_number": '',
    "w3_opening_song_title": '',
    "w3_opening_song_scripture": '',
    "w3_chairman": '',
    "w3_public_talk_start": '',
    "w3_public_talk_title": '',
    "w3_public_speaker": '',
    "w3_public_speaker_congregation": '',
    "w3_middle_song_start": '',
    "w3_middle_song_number": '',
    "w3_middle_song_title": '',
    "w3_middle_song_scripture": '',
    "w3_watchtower_start": '',
    "w3_watchtower_title": '',
    "w3_watchtower_conductor": '',
    "w3_watchtower_reader": '',
    "w3_closing_song_start": '',
    "w3_closing_song_number": '',
    "w3_closing_song_title": '',
    "w3_closing_song_scripture": '',
    "w3_closing_prayer": '',
    "w3_foyer_attendant": '',
    "w3_auditorium_attendant": '',
    "w3_stage_attendant": '',
    "w3_sound_console_operator": '',
    "w3_video_console_operator": '',
    "w3_mic_handler_1": '',
    "w3_mic_handler_2": '',
    "w3_mic_handler_3": '',
    "w3_mic_handler_4": '',
    "w3_mic_handler_5": '',
    "w3_mic_handler_6": '',
    "w3_hospitality_service_group": '',
    "w3_touch_up_cleaning_service_group": '',
    "w3_outgoing_speaker_1": '',
    "w3_outgoing_speaker_congregation_1": '',
    "w3_outgoing_speaker_2": '',
    "w3_outgoing_speaker_congregation_2": '',
    "w4_week": '',
    "w4_meeting_date": '',
    "w4_opening_song_start": '',
    "w4_opening_song_number": '',
    "w4_opening_song_title": '',
    "w4_opening_song_scripture": '',
    "w4_chairman": '',
    "w4_public_talk_start": '',
    "w4_public_talk_title": '',
    "w4_public_speaker": '',
    "w4_public_speaker_congregation": '',
    "w4_middle_song_start": '',
    "w4_middle_song_number": '',
    "w4_middle_song_title": '',
    "w4_middle_song_scripture": '',
    "w4_watchtower_start": '',
    "w4_watchtower_title": '',
    "w4_watchtower_conductor": '',
    "w4_watchtower_reader": '',
    "w4_closing_song_start": '',
    "w4_closing_song_number": '',
    "w4_closing_song_title": '',
    "w4_closing_song_scripture": '',
    "w4_closing_prayer": '',
    "w4_foyer_attendant": '',
    "w4_auditorium_attendant": '',
    "w4_stage_attendant": '',
    "w4_sound_console_operator": '',
    "w4_video_console_operator": '',
    "w4_mic_handler_1": '',
    "w4_mic_handler_2": '',
    "w4_mic_handler_3": '',
    "w4_mic_handler_4": '',
    "w4_mic_handler_5": '',
    "w4_mic_handler_6": '',
    "w4_hospitality_service_group": '',
    "w4_touch_up_cleaning_service_group": '',
    "w4_outgoing_speaker_1": '',
    "w4_outgoing_speaker_congregation_1": '',
    "w4_outgoing_speaker_2": '',
    "w4_outgoing_speaker_congregation_2": '',
    "w5_week": '',
    "w5_meeting_date": '',
    "w5_opening_song_start": '',
    "w5_opening_song_number": '',
    "w5_opening_song_title": '',
    "w5_opening_song_scripture": '',
    "w5_chairman": '',
    "w5_public_talk_start": '',
    "w5_public_talk_title": '',
    "w5_public_speaker": '',
    "w5_public_speaker_congregation": '',
    "w5_middle_song_start": '',
    "w5_middle_song_number": '',
    "w5_middle_song_title": '',
    "w5_middle_song_scripture": '',
    "w5_watchtower_start": '',
    "w5_watchtower_title": '',
    "w5_watchtower_conductor": '',
    "w5_watchtower_reader": '',
    "w5_closing_song_start": '',
    "w5_closing_song_number": '',
    "w5_closing_song_title": '',
    "w5_closing_song_scripture": '',
    "w5_closing_prayer": '',
    "w5_foyer_attendant": '',
    "w5_auditorium_attendant": '',
    "w5_stage_attendant": '',
    "w5_sound_console_operator": '',
    "w5_video_console_operator": '',
    "w5_mic_handler_1": '',
    "w5_mic_handler_2": '',
    "w5_mic_handler_3": '',
    "w5_mic_handler_4": '',
    "w5_mic_handler_5": '',
    "w5_mic_handler_6": '',
    "w5_hospitality_service_group": '',
    "w5_touch_up_cleaning_service_group": '',
    "w5_outgoing_speaker_1": '',
    "w5_outgoing_speaker_congregation_1": '',
    "w5_outgoing_speaker_2": '',
    "w5_outgoing_speaker_congregation_2": ''
}

# give bounds to template tables for each week
template_table_indexes = {
    "w1_start": 0,
    "w1_end": 7,
    "w2_start": 8,
    "w2_end": 15,
    "w3_start": 16,
    "w3_end": 23,
    "w4_start": 24,
    "w4_end": 31,
    "w5_start": 32,
    "w5_end": 39,
}

def render(meetings=[], docx_filename=docx_template_file):
    variables = dict(template_variables)
    number_of_meetings = len(meetings)
    for i in range(0, number_of_meetings):
        variables[f"w{i+1}_week"] = meetings[i]['week']
        variables[f"w{i+1}_meeting_date"] = meetings[i]['meeting_date']
        variables[f"w{i+1}_opening_song_start"] = meetings[i]['opening_song']['start']
        variables[f"w{i+1}_opening_song_number"] = meetings[i]['opening_song']['details']['number']
        variables[f"w{i+1}_opening_song_title"] = meetings[i]['opening_song']['details']['title']
        variables[f"w{i+1}_opening_song_scripture"] = meetings[i]['opening_song']['details']['ref_scripture']
        variables[f"w{i+1}_chairman"] = meetings[i]['chairman']
        variables[f"w{i+1}_public_talk_start"] = meetings[i]['public_talk']['start']
        variables[f"w{i+1}_public_title"] = meetings[i]['public_talk']['title']
        variables[f"w{i+1}_public_speaker"] = meetings[i]['public_talk']['speaker']
        variables[f"w{i+1}_public_speaker_congregation"] = meetings[i]['public_talk']['speaker_congregation']
        variables[f"w{i+1}_middle_song_start"] = meetings[i]['middle_song']['start']
        variables[f"w{i+1}_middle_song_number"] = meetings[i]['middle_song']['details']['number']
        variables[f"w{i+1}_middle_song_title"] = meetings[i]['middle_song']['details']['title']
        variables[f"w{i+1}_middle_song_scripture"] = meetings[i]['middle_song']['details']['ref_scripture']
        variables[f"w{i+1}_watchtower_start"] = meetings[i]['watchtower_study']['start']
        variables[f"w{i+1}_watchtower_title"] = meetings[i]['watchtower_study']['title']
        variables[f"w{i+1}_watchtower_conductor"] = meetings[i]['watchtower_study']['conductor']
        variables[f"w{i+1}_watchtower_reader"] = meetings[i]['watchtower_study']['reader']
        variables[f"w{i+1}_closing_song_start"] = meetings[i]['closing_song']['start']
        variables[f"w{i+1}_closing_song_number"] = meetings[i]['closing_song']['details']['number']
        variables[f"w{i+1}_closing_song_title"] = meetings[i]['closing_song']['details']['title']
        variables[f"w{i+1}_closing_song_scripture"] = meetings[i]['closing_song']['details']['ref_scripture']
        variables[f"w{i+1}_closing_prayer"] = None
        variables[f"w{i+1}_foyer_attendant"] = None
        variables[f"w{i+1}_auditorium_attendant"] = None
        variables[f"w{i+1}_sound_console_operator"] = None
        variables[f"w{i+1}_video_console_operator"] = None
        variables[f"w{i+1}_mic_handler_1"] = None
        variables[f"w{i+1}_mic_handler_2"] = None
        variables[f"w{i+1}_mic_handler_3"] = None
        variables[f"w{i+1}_mic_handler_4"] = None
        variables[f"w{i+1}_mic_handler_5"] = None
        variables[f"w{i+1}_mic_handler_6"] = None
        variables[f"w{i+1}_hospitality_service_group"] = None
        variables[f"w{i+1}_touch_up_cleaning_service_group"] = None
        variables[f"w{i+1}_outgoing_speaker_1"] = None
        variables[f"w{i+1}_outgoing_speaker_congregation_1"] = None
        variables[f"w{i+1}_outgoing_speaker_2"] = None
        variables[f"w{i+1}_outgoing_speaker_congregation_2"] = None
    if number_of_meetings < 4:
        variables["w4_week"] = '',
        variables["w4_opening_song_start"] = '',
        variables["w4_opening_song_numer"] = '',
        variables["w4_opening_song_title"] = '',
        variables["w4_opening_song_scripture"] = '',
        variables["w4_chairman"] = '',
        variables["w4_public_talk_start"] = '',
        variables["w4_public_talk_title"] = '',
        variables["w4_public_speaker"] = '',
        variables["w4_public_speaker_congregation"] = '',
        variables["w4_middle_song_start"] = '',
        variables["w4_middle_song_numer"] = '',
        variables["w4_middle_song_title"] = '',
        variables["w4_middle_song_scripture"] = '',
        variables["w4_watchtower_start"] = '',
        variables["w4_watchtower_title"] = '',
        variables["w4_watchtower_conductor"] = '',
        variables["w4_watchtower_reader"] = '',
        variables["w4_closing_song_start"] = '',
        variables["w4_closing_song_numer"] = '',
        variables["w4_closing_song_title"] = '',
        variables["w4_closing_song_scripture"] = '',
        variables["w4_closing_prayer"] = '',
        variables["w4_foyer_attendant"] = '',
        variables["w4_auditorium_attendant"] = '',
        variables["w4_stage_attendant"] = '',
        variables["w4_sound_console_operator"] = '',
        variables["w4_video_console_operator"] = '',
        variables["w4_mic_handler_1"] = '',
        variables["w4_mic_handler_2"] = '',
        variables["w4_mic_handler_3"] = '',
        variables["w4_mic_handler_4"] = '',
        variables["w4_mic_handler_5"] = '',
        variables["w4_mic_handler_6"] = '',
        variables["w4_hospitality_service_group"] = '',
        variables["w4_touch_up_cleaning_service_group"] = '',
        variables["w4_outgoing_speaker_1"] = '',
        variables["w4_outgoing_speaker_congregation_1"] = '',
        variables["w4_outgoing_speaker_2"] = '',
        variables["w4_outgoing_speaker_congregation_2"] = ''
    if number_of_meetings < 5:
        variables["w5_week"] = '',
        variables["w5_opening_song_start"] = '',
        variables["w5_opening_song_numer"] = '',
        variables["w5_opening_song_title"] = '',
        variables["w5_opening_song_scripture"] = '',
        variables["w5_chairman"] = '',
        variables["w5_public_talk_start"] = '',
        variables["w5_public_talk_title"] = '',
        variables["w5_public_speaker"] = '',
        variables["w5_public_speaker_congregation"] = '',
        variables["w5_middle_song_start"] = '',
        variables["w5_middle_song_numer"] = '',
        variables["w5_middle_song_title"] = '',
        variables["w5_middle_song_scripture"] = '',
        variables["w5_watchtower_start"] = '',
        variables["w5_watchtower_title"] = '',
        variables["w5_watchtower_conductor"] = '',
        variables["w5_watchtower_reader"] = '',
        variables["w5_closing_song_start"] = '',
        variables["w5_closing_song_numer"] = '',
        variables["w5_closing_song_title"] = '',
        variables["w5_closing_song_scripture"] = '',
        variables["w5_closing_prayer"] = '',
        variables["w5_foyer_attendant"] = '',
        variables["w5_auditorium_attendant"] = '',
        variables["w5_stage_attendant"] = '',
        variables["w5_sound_console_operator"] = '',
        variables["w5_video_console_operator"] = '',
        variables["w5_mic_handler_1"] = '',
        variables["w5_mic_handler_2"] = '',
        variables["w5_mic_handler_3"] = '',
        variables["w5_mic_handler_4"] = '',
        variables["w5_mic_handler_5"] = '',
        variables["w5_mic_handler_6"] = '',
        variables["w5_hospitality_service_group"] = '',
        variables["w5_touch_up_cleaning_service_group"] = '',
        variables["w5_outgoing_speaker_1"] = '',
        variables["w5_outgoing_speaker_congregation_1"] = '',
        variables["w5_outgoing_speaker_2"] = '',
        variables["w5_outgoing_speaker_congregation_2"] = ''
    dir_path = os.path.dirname(os.path.realpath(__file__))
    doc = Document(docx_filename)
    var_regex = re.compile("{{.*}}")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        replacement_text = str(r.text)
                        found_variable = False
                        for m in re.finditer(r'\{\{[^\}\}]*\}\}', r.text):
                            replace_str = r.text[m.span()[0]:m.span()[1]]
                            variable_name = r.text[m.span()[0]+3:m.span()[1]-3].strip()
                            variable_value = ''
                            if variable_name in variables:
                                variable_value = variables[variable_name]
                                found_variable = True
                            replacement_text = replacement_text.replace(replace_str, str(variable_value))
                        if found_variable:
                            r.text = replacement_text
    # remove unused template tables
    for i, t in enumerate(doc.tables):
        if number_of_meetings < 4 and i >= template_table_indexes['w4_start']:
            table_element = t._element
            part_element = table_element.getnext()
            body_element = table_element.getparent()
            body_element.remove(table_element)
            body_element.remove(part_element)
        elif number_of_meetings < 5 and i >= template_table_indexes['w5_start']:
            table_element = t._element
            part_element = table_element.getnext()
            body_element = t._element.getparent()
            body_element.remove(table_element)
            body_element.remove(part_element)
    # remove unused page breaks
    if number_of_meetings < 4:
        last_table_element = doc.tables[template_table_indexes['w3_end']]._element
        paragraph_element = last_table_element.getnext()
        for child in paragraph_element.getchildren():
            if isinstance(child, CT_R):
                paragraph_element.remove(child)
    if number_of_meetings < 5:
        last_table_element = doc.tables[template_table_indexes['w4_end']]._element
        paragraph_element = last_table_element.getnext()
        for child in paragraph_element.getchildren():
            if isinstance(child, CT_R):
                paragraph_element.remove(child)
    return doc
